from io import BytesIO

from docx import Document

from app.resume.parser import parse_docx_bytes


def test_parse_docx_bytes_extracts_text() -> None:
    buf = BytesIO()
    doc = Document()
    doc.add_heading("Jane Doe", level=1)
    doc.add_paragraph("Built a FastAPI backend.")
    doc.save(buf)

    text = parse_docx_bytes(buf.getvalue())
    assert "Jane Doe" in text
    assert "FastAPI backend" in text

