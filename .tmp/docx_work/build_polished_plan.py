from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "doubow_build_plan_polished.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "6B7280"
HEADER_FILL = "F2F4F7"
SOFT_FILL = "F4F6F9"
BORDER = "D9DEE7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, col_widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(col_widths)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in col_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(col_widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(col_widths[idx]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Title", 26, "0B2545", 0, 8),
        ("Subtitle", 11, MUTED, 0, 18),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name in {"Title", "Heading 1", "Heading 2", "Heading 3"}
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_para(doc, text="", style=None, bold_label=None):
    p = doc.add_paragraph(style=style)
    if bold_label and text.startswith(bold_label):
        r = p.add_run(bold_label)
        set_font(r, bold=True)
        r2 = p.add_run(text[len(bold_label):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr[idx].text = ""
        p = hdr[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, bold=True, color=INK)
        set_cell_shading(hdr[idx], HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r)
    return table


def add_gate(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    set_table_borders(table, color="D5DAE3")
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, SOFT_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Gate: ")
    set_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_font(r)
    return table


def add_phase(doc, title, staffing, deliverables, gate):
    add_para(doc, title, style="Heading 2")
    p = add_para(doc, staffing)
    p.runs[0].italic = True
    for item in deliverables:
        add_bullet(doc, item)
    add_gate(doc, gate)


doc = Document()
style_doc(doc)

title = doc.add_paragraph(style="Title")
run = title.add_run("Doubow")
set_font(run, size=26, bold=True, color="0B2545")
subtitle = doc.add_paragraph(style="Subtitle")
run = subtitle.add_run("Engineering Build Plan | Version 1.0 | May 2026 | 16-week MVP plan")
set_font(run, size=11, color=MUTED)

add_para(
    doc,
    "This plan defines the technical architecture, delivery phases, dependencies, risks, staffing model, and MVP acceptance criteria for the Doubow build. The emphasis is on a controlled human-in-the-loop application pipeline, reliable AI output validation, and production readiness before launch.",
)

add_para(doc, "Tech Stack", style="Heading 1")
add_table(
    doc,
    ["Layer", "Technology", "Rationale"],
    [
        ("Frontend", "Next.js 14 + TypeScript", "Server rendering/static generation, file-based routing, and a mature application ecosystem."),
        ("UI components", "shadcn/ui + Tailwind CSS", "Accessible primitives with fast iteration and a flexible design system."),
        ("Backend API", "FastAPI on Python 3.12", "Async-native API development, Pydantic validation, and generated OpenAPI documentation."),
        ("Task queue", "Celery + Redis", "Background processing for scraping, AI workflows, PDF parsing, and notifications."),
        ("Database", "PostgreSQL 16 + pgvector", "Relational workflow state plus vector similarity search."),
        ("AI / agents", "LangChain, modular architecture", "Composable agent and tool orchestration."),
        ("LLM providers", "Anthropic Claude primary; OpenAI fallback", "Quality-focused primary provider with redundancy for availability."),
        ("Object storage", "AWS S3 or compatible storage", "Storage for resume files, raw HTML, and generated PDFs."),
        ("Authentication", "NextAuth.js + JWT", "Session management and OAuth support for Google and LinkedIn."),
        ("Payments", "Stripe + Customer Portal", "Subscription billing, customer self-service, and webhook-driven account state."),
        ("Observability", "Prometheus + Grafana + Sentry", "Operational metrics, dashboards, and error tracking."),
        ("CI/CD", "GitHub Actions -> Docker -> ECS", "Reproducible builds and zero-downtime deployment workflow."),
    ],
    [1900, 2700, 4760],
)

add_para(doc, "Application State Machine", style="Heading 1")
add_para(doc, "Canonical flow: DISCOVERED -> SCORING -> DRAFTING -> PENDING_APPROVAL -> APPROVED -> SUBMITTED. Any state may transition to FAILED, then RETRY resumes from the failed step.")
add_gate(doc, "The /applications/{id}/submit endpoint must verify status = APPROVED in the database before dispatch. If the status is not APPROVED, the API returns HTTP 403; this authorization check belongs in the API handler and must be covered by penetration testing.")

add_para(doc, "Engineering Phases", style="Heading 1")
add_phase(
    doc,
    "Phase 1: Foundation & Plumbing",
    "Weeks 1-3 | 2 backend engineers + 1 DevOps engineer",
    [
        "Create the Postgres schema for users, profiles, jobs, applications, outreach_drafts, llm_logs, check_ins, and milestones.",
        "Implement the application state machine with Alembic migrations.",
        "Establish the FastAPI project structure with Pydantic v2 schemas.",
        "Configure Celery + Redis queues for scrape, score, draft, and notify jobs, including a dead-letter queue.",
        "Implement JWT authentication plus Google and LinkedIn OAuth.",
        "Build the resume upload worker: PDF/DOCX parsing -> structured JSON -> text-embedding-3-small embedding.",
        "Add idempotency-key middleware that deduplicates requests within a 24-hour window.",
        "Configure S3 object storage for raw resume files.",
    ],
    "State-machine transitions are enforced. POST /applications/{id}/submit returns 403 unless the database status is APPROVED.",
)
add_phase(
    doc,
    "Phase 2: RAG & Matching Core",
    "Weeks 4-6 | 2 backend engineers + 1 ML engineer",
    [
        "Build pluggable, rate-limited job scraper workers by source; store raw HTML in S3.",
        "Create the job embedding pipeline with text-embedding-3-small and pgvector HNSW indexing.",
        "Implement the Fit Scorer agent with structured JSON output: score, match_pct, rationale, gap_skills, and strength_skills.",
        "Validate every agent output with Pydantic; invalid results move to FAILED and trigger an alert.",
        "Ship the job catalog REST API with filtering, sorting, and pagination.",
        "Create the match-feed endpoint using cosine-similarity ranking against the user's resume embedding.",
        "Deduplicate jobs with hash(source_url) as the unique key per source.",
    ],
    "The scorer produces reliable structured JSON, and invalid outputs are caught before they reach the database.",
)
add_phase(
    doc,
    "Phase 3: Agentic Layer",
    "Weeks 7-9 | 2 backend engineers + 1 ML engineer",
    [
        "Build the Drafter Agent; store email and LinkedIn message drafts in outreach_drafts with status = DRAFT.",
        "Build the Interview Prep Agent with RAG grounding from the job description and resume context.",
        "Build the Sender Agent so it can dispatch email or LinkedIn messages only when applications.status = APPROVED.",
        "Log LLM interactions: agent_name, prompt_hash, raw_output, latency_ms, user_edit, and feedback_score.",
        "Handle RETRY and FAILED states with exponential backoff, capped at three retries.",
        "Build the unified Career Copilot as a LangChain agent with six structured tools and WebSocket streaming.",
        "Store conversation history per session, using the last 20 turns as context.",
    ],
    "The DISCOVERED -> PENDING_APPROVAL flow works end to end, and the Sender cannot fire without APPROVED in the database.",
)
add_phase(
    doc,
    "Phase 4: UI & Approval Dashboard",
    "Weeks 10-13 | 2 frontend engineers + 1 backend engineer",
    [
        "Ship P0 MVP-critical pages: Onboarding wizard, Career Profile, Dashboard, Approval Dashboard, Job Discovery, Job Tracker, and Career Copilot.",
        "Prepare P1 launch-complete pages: Career Planner, Career Pathfinder, Career Success, ATS Optimizer, and Settings/Stripe.",
        "Backlog P2 post-launch surfaces: CV Builder, Cover Letter, Career Health, LinkedIn Analysis, Salary Benchmark, Sponsorship Hub, and Discussion Board.",
        "Implement the design system: dark navy sidebar (#0F1117), white content area, blue accent (#4F8EF7), and shadcn/ui components.",
        "Use TanStack Query for server state and Zustand for client state.",
        "Build the Approval Dashboard with inline editing, approve -> APPROVED, reject, and real-time WebSocket status updates.",
    ],
    "All P0 pages are shipped and QA-approved. The approval flow is tested end to end with human-in-the-loop enforcement verified.",
)
add_phase(
    doc,
    "Phase 5: Production Hardening",
    "Weeks 14-16 | Full team",
    [
        "Instrument Prometheus metrics for request latency, worker queue depth, FAILED rate, LLM latency, and scraper success.",
        "Create Grafana dashboards for pipeline health, agent performance, and the user funnel.",
        "Apply rate limits for scrapers by domain, LLM usage by token budget per tier, and API access at 100 requests per minute per user.",
        "Implement GDPR support: /users/me/delete endpoint, audit-log table, data-retention policy, and cookie consent.",
        "Add monthly drift-detection cron checks that alert when agent edit rate exceeds 40%.",
        "Run k6 load tests for 500 concurrent users with p95 latency below 500 ms.",
        "Finalize GitHub Actions CI/CD: lint -> test -> build -> push -> deploy -> smoke test -> promote to production.",
    ],
    "Production telemetry, compliance controls, load testing, and deployment automation are in place for MVP launch.",
)

add_para(doc, "Dependency Map", style="Heading 1")
add_para(doc, "External dependencies to acquire in Week 1:", style="Heading 2")
for dep in [
    "Anthropic API key for the primary LLM provider.",
    "OpenAI API key for embeddings and fallback LLM coverage.",
    "LinkedIn API access; allow 4-6 weeks for approval because this may sit on the critical path.",
    "AWS S3 bucket and IAM roles.",
    "Stripe account and webhook endpoints.",
    "Email sending service, either SES or Resend.",
]:
    add_bullet(doc, dep)

add_para(doc, "Risk Register", style="Heading 1")
add_table(
    doc,
    ["Risk", "Severity", "Mitigation"],
    [
        ("HITL enforcement bypassed through direct API access", "Critical", "Return 403 from the Sender path unless database status is APPROVED; verify with security tests."),
        ("Scrapers blocked by CAPTCHA, rate limits, or DOM changes", "Critical", "Decouple scraping by provider, store raw HTML separately, and keep downstream workflows tolerant of scraper failure."),
        ("LinkedIn API access rejected or revoked", "Critical", "Keep Sender provider-agnostic and maintain an email-first fallback."),
        ("LLM output schema drift", "High", "Validate all agent outputs with Pydantic; invalid results move to FAILED and trigger alerts."),
        ("Resume parsing quality degrades scoring", "High", "Allow users to edit parsed fields, add extraction quality checks, and provide manual-entry fallback."),
        ("Celery task duplication after restart", "High", "Use idempotency keys for every sensitive task, delivered in Phase 1."),
        ("UI scope underestimated", "Medium", "Launch P0 pages only; move P1/P2 scope into phased rollout."),
        ("GDPR / PII compliance gaps", "Medium", "Design schema-level privacy and deletion controls in Phase 1 rather than retrofitting them later."),
        ("Scoring drift over time", "Low", "Run monthly drift detection comparing user edits against model recommendations."),
    ],
    [3350, 1350, 4660],
)

add_para(doc, "Team & Responsibilities", style="Heading 1")
add_table(
    doc,
    ["Role", "Count", "Responsibilities"],
    [
        ("Backend Engineer", "2", "FastAPI, state machine, database schema, workers, and agent integration."),
        ("ML / AI Engineer", "1", "Agent design, RAG pipeline, embedding quality, and prompt engineering."),
        ("Frontend Engineer", "2", "Next.js UI, component library, approval dashboard, and chat interface."),
        ("DevOps Engineer", "1", "Infrastructure, CI/CD, observability, and security hardening."),
        ("Product Manager", "1", "Requirements, prioritization, stakeholder communication, and KPI tracking."),
        ("Designer", "1", "Design system, component specifications, user testing, and accessibility."),
    ],
    [2600, 900, 5860],
)

add_para(doc, "Definition of Done for MVP", style="Heading 1")
for item in [
    "Users can register, upload a resume, and complete the Career Profile.",
    "Job Discovery displays a personalized match feed with fit scores.",
    "Users can track applications through the full pipeline.",
    "AI generates outreach drafts and surfaces them in the Approval Dashboard.",
    "No draft can be sent without explicit approval, verified by penetration testing.",
    "Career Copilot answers career questions using context from the user's resume.",
    "ATS Optimizer returns keyword-match analysis and improvement suggestions.",
    "Stripe subscription flow works end to end for Standard, Pro, and Ultimate tiers.",
    "All P0 pages pass a WCAG 2.1 AA audit.",
    "API p95 response time is below 300 ms under 500 concurrent users.",
    "GDPR deletion endpoint is tested and verified.",
]:
    add_bullet(doc, item)

footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = footer.add_run("Doubow Engineering Build Plan")
set_font(run, size=9, color=MUTED)

doc.save(OUT)
